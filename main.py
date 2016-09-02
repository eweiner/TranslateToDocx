from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from DatabaseScraper import DBSScraper

def write_table_to_document(content, document, table, tableName, pathSettings=False, securityenvs=False):
    schema = table[tableName]
    document.add_paragraph(tableName, style = "Heading 1")
    keys = None
    for i in range(0, len(content)):
        document.add_paragraph(content[i][1], style = "Heading 2")
        for j in range(0, len(content[i])):
            if(type(content[i][j]) is dict):
                if(pathSettings):
                    keys = content[i][j][content[i][0]].keys()

                    e = content
                    document.add_paragraph(schema[j] + ":", style = "List Bullet")
                    for key in keys:
                        document.add_paragraph(str(key) + ": " + str(content[i][j][content[i][0]][key]), style = "List Bullet 2")
                if(securityenvs):
                    keyKeys = content[i][j]["keys"].keys()
                    algKeys = content[i][j]["algs"].keys()

            else:
                document.add_paragraph(schema[j] + ": " + str(content[i][j]), style = "List Bullet")


def initializeDoc(path):
    document = Document()
    styles = document.styles()
    return styles, document

def createStyle(style, wd_style, fontName, size, name = 'Times New Roman', color=[0, 0, 0], biu = [False, False, False]):
    charstyle = style.add_style(fontName, wd_style)
    font = charstyle.font
    font.name = name
    font.size = Pt(size)
    font.bold = biu[0]
    font.italic = biu[1]
    font.underline = biu[2]
    font.color.rgb = RGBColor(color[0], color[1], color[2])

#Iterate through sqlite_master which has a lot of extraneous data in order to find the schema of a table
#Returns a dictionary with a key as the name of the category and the contents the variable names


def main():
    db_file = "/Users/ericweiner/Downloads/TestSDB.sdb"
    scraper = DBSScraper(db_file)
    table = scraper.get_table()
    resources = scraper.get_resources()
    path_settings = scraper.get_path_settings()
    security_environments = scraper.get_security_envs()
    document = Document()
    document.add_heading("Database Settings", 0)
    write_table_to_document(resources, document, table, "Resources")
    write_table_to_document(path_settings, document, table, "PathSettings", pathSettings=True)
    write_table_to_document(security_environments, document, table, "SecurityEnvironments")
    document.save("/Users/ericweiner/Documents/test.docx")


main()